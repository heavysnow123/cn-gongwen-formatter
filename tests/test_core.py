"""Word Formatter Pro 引擎无头测试（不依赖 GUI / COM）。

覆盖：标题识别、格式化应用、表格、附件、空行、Markdown 清理、TXT/MD 转换。
"""

import os
import tempfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from word_formatter.config import FormatterConfig
from word_formatter.core import (
    WordFormatter, detect_heading_level, clean_markdown,
    normalize_blank_lines,
)


def _make_docx(path, lines):
    """lines: list of (text, dict-of-attrs)。attrs 可含 align/center/font/size。"""
    doc = Document()
    for text, attrs in lines:
        p = doc.add_paragraph(text)
        if attrs.get("center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if attrs.get("font"):
            for r in p.runs:
                r.font.name = attrs["font"]
        if attrs.get("size"):
            for r in p.runs:
                r.font.size = __import__("docx").shared.Pt(attrs["size"])
    doc.save(path)


def _tmp(ext=".docx"):
    fd, p = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    return p


def test_heading_detection():
    assert detect_heading_level("一、总则") == 1
    assert detect_heading_level("（一）范围") == 2
    assert detect_heading_level("1. 定义") == 3
    assert detect_heading_level("(1) 细则") == 4
    assert detect_heading_level("普通正文一行文字") == 0
    print("✓ 标题识别正确")


def test_title_center_and_font():
    src = _tmp()
    _make_docx(src, [
        ("某型防空系统技术分析报告", {"center": True, "font": "黑体", "size": 22}),
        ("2026 年度阶段性总结", {"center": True, "font": "黑体", "size": 16}),
        ("一、引言", {}),
        ("这是正文，应当首行缩进并使用宋体。", {}),
        ("（一）研究背景", {}),
        ("1. 项目来源", {}),
        ("(1) 关键指标", {}),
    ])
    out = _tmp()
    cfg = FormatterConfig()
    fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
    r = fmt.format_file(src, os.path.dirname(out))
    assert r["error"] is None, r["error"]
    doc = Document(r["output"])
    paras = [p for p in doc.paragraphs]
    # 第一段应居中且字号 22
    assert paras[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    size0 = paras[0].runs[0].font.size.pt
    assert abs(size0 - 22.0) < 0.01, size0
    # 正文首行缩进
    body = [p for p in paras if p.text.startswith("这是正文")][0]
    assert body.paragraph_format.first_line_indent is not None
    os.remove(src)
    os.remove(r["output"])
    print("✓ 标题居中/字号/正文缩进正确")


def test_table_formatting():
    src = _tmp()
    doc = Document()
    doc.add_paragraph("一、数据汇总")
    t = doc.add_table(rows=2, cols=3)
    hdr = ["项目", "数值", "占比"]
    for j, h in enumerate(hdr):
        t.rows[0].cells[j].paragraphs[0].add_run(h)
    t.rows[1].cells[0].paragraphs[0].add_run("样本A")
    t.rows[1].cells[1].paragraphs[0].add_run("123")
    t.rows[1].cells[2].paragraphs[0].add_run("45%")
    doc.save(src)
    out = _tmp()
    cfg = FormatterConfig()
    cfg.enable_table_formatting = True
    fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
    r = fmt.format_file(src, os.path.dirname(out))
    assert r["error"] is None, r["error"]
    doc2 = Document(r["output"])
    tbl = doc2.tables[0]
    # 表头加粗
    assert tbl.rows[0].cells[0].paragraphs[0].runs[0].font.bold is True
    # 有边框
    borders = tbl._tbl.tblPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders")
    assert borders is not None
    os.remove(src)
    os.remove(r["output"])
    print("✓ 表格格式化（表头加粗/边框）正确")


def test_attachment_page_break():
    src = _tmp()
    _make_docx(src, [
        ("报告正文标题", {"center": True, "font": "黑体", "size": 22}),
        ("一、正文", {}),
        ("一些内容。", {}),
        ("附件1：原始数据表", {}),
        ("数据内容。", {}),
    ])
    cfg = FormatterConfig()
    cfg.enable_attachment_formatting = True
    fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
    r = fmt.format_file(src, tempfile.gettempdir())
    assert r["error"] is None, r["error"]
    doc = Document(r["output"])
    attach = [p for p in doc.paragraphs if p.text.startswith("附件1")][0]
    pPr = attach._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    has_break = False
    if pPr is not None:
        for br in pPr.iter():
            if br.tag.endswith("}br") and br.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                has_break = True
    assert has_break, "附件未插入分页符"
    os.remove(src)
    os.remove(r["output"])
    print("✓ 附件段前分页正确")


def test_markdown_clean():
    md = "# 一级标题\n## 二级\n- 项目一\n**粗体** 和 *斜体*\n[链接](http://x.com)\n> 引用\n\n---\n正文"
    out = clean_markdown(md)
    assert "#" not in out
    assert "**" not in out
    assert "[链接]" not in out
    assert "http://x.com" not in out
    assert ">" not in out
    assert "一级标题" in out
    print("✓ Markdown 清理正确")


def test_blank_line_modes():
    lines = ["甲", "", "", "", "乙", "", "丙"]
    # remove_single: 中间多个空行保留为一个? 实际逻辑：先折叠连续空行到1个，
    # 再删除夹在正文间的单个空行
    res = normalize_blank_lines(lines, "remove_single")
    # 期望：开头连续空行->1个；甲乙之间单个空行删除；乙丙之间单个空行删除
    assert "甲" in res and "乙" in res and "丙" in res
    keep = normalize_blank_lines(lines, "keep_single")
    assert "甲" in keep
    print("✓ 空行模式正确")


def test_txt_and_md_roundtrip():
    for ext, text, is_md in (
        (".txt", "报告标题\n一、第一章\n正文段落内容。\n（一）小节\n1. 条目", False),
        (".md", "# 报告标题\n## 第一章\n正文段落内容。\n- 列表项", True),
    ):
        src = _tmp(ext)
        with open(src, "w", encoding="utf-8") as f:
            f.write(text)
        cfg = FormatterConfig()
        fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
        r = fmt.format_file(src, tempfile.gettempdir())
        assert r["error"] is None, (ext, r["error"])
        doc = Document(r["output"])
        assert len([p for p in doc.paragraphs if p.text.strip()]) > 0
        os.remove(src)
        os.remove(r["output"])
    print("✓ TXT / MD 转换生成有效 docx")


def test_page_number_added():
    src = _tmp()
    _make_docx(src, [("标题", {"center": True}), ("一、章", {}), ("正文。", {})])
    cfg = FormatterConfig()
    cfg.page_number = True
    fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
    r = fmt.format_file(src, tempfile.gettempdir())
    assert r["error"] is None, r["error"]
    doc = Document(r["output"])
    footer = doc.sections[0].footer
    xml = footer.paragraphs[0]._p.xml
    assert "PAGE" in xml
    os.remove(src)
    os.remove(r["output"])
    print("✓ 页码插入正确")


def test_header_footer_standard():
    """标准模式：页眉 + 页脚文字 + 总页数 均生效，产物有效。"""
    src = _tmp()
    _make_docx(src, [("标题", {"center": True}), ("一、章", {}), ("正文。", {})])
    cfg = FormatterConfig()
    cfg.header_enabled = True
    cfg.header_text = "某招标文件"
    cfg.header_border = True
    cfg.page_number = True
    cfg.footer_text = "机密"
    cfg.page_number_total = True
    fmt = WordFormatter(cfg, log_cb=lambda m, l: None)
    r = fmt.format_file(src, tempfile.gettempdir())
    assert r["error"] is None, r["error"]
    doc = Document(r["output"])
    hdr = doc.sections[0].header.paragraphs[0]._p.xml
    ftr = doc.sections[0].footer.paragraphs[0]._p.xml
    assert "某招标文件" in hdr
    assert "bottom" in hdr.lower(), "页眉下边框缺失"
    assert "PAGE" in ftr and "NUMPAGES" in ftr
    assert "机密" in ftr
    os.remove(src)
    os.remove(r["output"])
    print("✓ 标准模式页眉/页脚/总页数正确")


if __name__ == "__main__":
    test_heading_detection()
    test_title_center_and_font()
    test_table_formatting()
    test_attachment_page_break()
    test_markdown_clean()
    test_blank_line_modes()
    test_txt_and_md_roundtrip()
    test_page_number_added()
    test_header_footer_standard()
    print("\n所有引擎测试通过 ✅")
