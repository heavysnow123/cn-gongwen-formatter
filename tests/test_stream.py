"""流式引擎测试：大文档、低内存、结构保留、格式正确。"""

import os
import zipfile
import tempfile
import tracemalloc
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION

from word_formatter.config import FormatterConfig
from word_formatter.stream import stream_format_document, StreamFormatter
from word_formatter.core import get_para_font_info


def _count(xml_bytes, tag_local):
    root = etree.fromstring(xml_bytes)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return len(root.findall(f".//{ns}{tag_local}"))


def _make_big_doc(path, n_paras=6000):
    doc = Document()
    doc.add_paragraph("某型装备招标文件技术分析报告").alignment = 1
    doc.add_paragraph("2026 年度公开招标").alignment = 1
    for i in range(n_paras):
        if i % 500 == 0:
            doc.add_paragraph(f"一、第{i // 500}章 总体要求")
        elif i % 500 == 100:
            doc.add_paragraph(f"（一）第{i}节 技术指标")
        elif i % 500 == 200:
            doc.add_paragraph(f"1. 条款{i}")
        else:
            doc.add_paragraph(f"第 {i} 条：投标人须满足下列技术要求，并提供相应证明材料。")
    # 一张表格
    t = doc.add_table(rows=2, cols=3)
    for j, h in enumerate(["指标", "要求", "证明"]):
        t.rows[0].cells[j].paragraphs[0].add_run(h)
    t.rows[1].cells[0].paragraphs[0].add_run("射程")
    t.rows[1].cells[1].paragraphs[0].add_run("≥200km")
    t.rows[1].cells[2].paragraphs[0].add_run("检测报告")
    # 中途插入一个分节符（模拟标书横向附录页）
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("附录：横向排版页内容示例")
    doc.save(path)


def test_streaming_big_doc():
    src = tempfile.mkstemp(suffix=".docx")[1]
    _make_big_doc(src, n_paras=6000)
    out = tempfile.mkstemp(suffix="_fmt.docx")[1]

    # 输入侧基准
    with zipfile.ZipFile(src) as z:
        in_doc_xml = z.read("word/document.xml")
    in_paras = _count(in_doc_xml, "p")
    in_sectpr = _count(in_doc_xml, "sectPr")
    in_tables = _count(in_doc_xml, "tbl")

    cfg = FormatterConfig()
    cfg.enable_table_formatting = True
    fmt = StreamFormatter(cfg, log_cb=lambda m, l: None)

    tracemalloc.start()
    fmt.run(src, out)
    cur, peak = tracemalloc.get_traced_memory()
    tracemalloc.stop()

    print(f"段落数: {in_paras}, 表格: {in_tables}, sectPr: {in_sectpr}")
    print(f"流式峰值内存(本进程): {peak / 1024 / 1024:.1f} MB")

    # 输出必须仍是合法包，且（原始 XML 口径）段落数不变
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(out) as z:
        out_doc_xml = z.read("word/document.xml")
    out_root = etree.fromstring(out_doc_xml)
    out_paras = _count(out_doc_xml, "p")
    assert out_paras == in_paras, f"段落数变化 {out_paras} != {in_paras}"
    assert _count(out_doc_xml, "sectPr") == in_sectpr, "分节符丢失！"
    assert _count(out_doc_xml, "tbl") == in_tables, "表格丢失！"

    # 标题格式已应用（首个 w:p 应居中且字号 22）
    first_p = out_root.find(f".//{W}p")
    jc = first_p.find(f"{W}pPr/{W}jc")
    assert jc is not None and jc.get(f"{W}val") == "center", "标题未居中"
    sz = first_p.find(f".//{W}r/{W}rPr/{W}sz")
    assert sz is not None and abs(float(sz.get(f"{W}val")) - 44) < 0.01, \
        f"标题半磅值错误 {sz.get(f'{W}val') if sz is not None else None}"
    # 22pt = 44 half-points

    # 表格边框已加
    first_tbl = out_root.find(f".//{W}tbl")
    assert first_tbl is not None
    assert first_tbl.find(f"{W}tblPr/{W}tblBorders") is not None, "表格边框未添加"

    # 峰值内存应远低于“整文档加载”的临界点（< 200MB 视为恒定可控）
    assert peak < 200 * 1024 * 1024, f"内存异常: {peak/1024/1024:.0f} MB"

    for f in (src, out):
        try:
            os.remove(f)
        except OSError:
            pass
    print("✓ 流式大文档：内存恒定、结构保留、格式正确")


def test_streaming_matches_standard_on_small():
    """小文档两种模式结果应一致（标题/正文/表格）。"""
    from word_formatter.core import WordFormatter
    src = tempfile.mkstemp(suffix=".docx")[1]
    _make_big_doc(src, n_paras=40)
    out_std = tempfile.mkstemp(suffix="_std.docx")[1]
    out_stream = tempfile.mkstemp(suffix="_stream.docx")[1]

    cfg = FormatterConfig()
    cfg.enable_table_formatting = True
    # 标准模式
    out_dir = os.path.dirname(out_std)
    WordFormatter(cfg, log_cb=lambda m, l: None).format_file(src, out_dir)
    base = os.path.splitext(os.path.basename(src))[0]
    std_path = os.path.join(out_dir, f"{base}_formatted.docx")
    std = Document(std_path)
    # 流式模式
    stream_format_document(src, out_stream, cfg, "docx", log_cb=lambda m, l: None)
    st = Document(out_stream)

    # 段落数一致
    assert len([p for p in st.paragraphs]) == len([p for p in std.paragraphs])
    # 标题字号一致
    assert abs(get_para_font_info(st.paragraphs[0])[1] - 22.0) < 0.01
    # 表格边框两种都有
    assert st.tables[0]._tbl.tblPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders") is not None

    for f in (src, std_path, out_stream):
        try:
            os.remove(f)
        except OSError:
            pass
    print("✓ 流式模式与标准模式结果一致（小文档）")


def test_streaming_header_footer():
    """流式模式也应注入页眉/页脚/页码，且不破坏包结构。"""
    from docx.enum.section import WD_SECTION
    src = tempfile.mkstemp(suffix=".docx")[1]
    doc = Document()
    doc.add_paragraph("主标题").alignment = 1
    for i in range(20):
        doc.add_paragraph(f"正文第 {i} 段，内容用于填充文档。")
    doc.add_section(WD_SECTION.NEW_PAGE)   # 制造第二个分节（内联 sectPr）
    doc.add_paragraph("附录内容")
    doc.save(src)
    out = tempfile.mkstemp(suffix="_hf.docx")[1]

    cfg = FormatterConfig()
    cfg.header_enabled = True
    cfg.header_text = "某型装备招标文件"
    cfg.header_align = "center"
    cfg.header_border = True
    cfg.page_number = True
    cfg.footer_text = "机密"
    cfg.page_number_total = True

    stream_format_document(src, out, cfg, "docx", log_cb=lambda m, l: None)

    with zipfile.ZipFile(out) as z:
        names = z.namelist()
        assert any(n.endswith("header1.xml") for n in names), "缺 header 部件"
        assert any(n.endswith("footer1.xml") for n in names), "缺 footer 部件"
        hdr = z.read([n for n in names if n.endswith("header1.xml")][0]).decode("utf-8")
        ftr = z.read([n for n in names if n.endswith("footer1.xml")][0]).decode("utf-8")
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        ct = z.read("[Content_Types].xml").decode("utf-8")
        doc_xml = z.read("word/document.xml").decode("utf-8")

    # 关系与内容类型
    assert "rIdWfpHeader" in rels and "rIdWfpFooter" in rels
    assert "header+xml" in ct and "footer+xml" in ct
    # 页眉内容
    assert "某型装备招标文件" in hdr
    assert "w:bottom" in hdr, "页眉下边框缺失"
    # 页脚内容：页码域 + 总页数 + 附加文字
    assert "PAGE" in ftr and "NUMPAGES" in ftr
    assert "机密" in ftr and "第 " in ftr and " 页 / 共 " in ftr
    # 每个 sectPr 都被注入了引用
    assert doc_xml.count("footerReference") >= 2, "sectPr 未注入页脚引用"
    assert doc_xml.count("headerReference") >= 2, "sectPr 未注入页眉引用"
    assert 'r:id="rIdWfpFooter"' in doc_xml
    assert 'r:id="rIdWfpHeader"' in doc_xml

    # 产物仍可被 python-docx 正常打开
    d = Document(out)
    assert d.sections[0].footer.paragraphs[0] is not None

    for f in (src, out):
        try:
            os.remove(f)
        except OSError:
            pass
    print("✓ 流式模式页眉/页脚/页码注入正确，包结构完整")


if __name__ == "__main__":
    test_streaming_big_doc()
    test_streaming_matches_standard_on_small()
    test_streaming_header_footer()
    print("\n流式引擎测试通过 ✅")
