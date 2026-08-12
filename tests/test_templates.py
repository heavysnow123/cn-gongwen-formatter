import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from docx import Document
from docx.shared import RGBColor
from word_formatter.templates import (
    gongwen_template, report_template,
    generate_gongwen, generate_report, generate_redhead,
    TEMPLATES, TEMPLATE_DESCRIPTIONS, GENERATORS, generate_template,
)
from word_formatter.config import FormatterConfig


def test_gongwen_config():
    c = gongwen_template()
    assert c.body_font == "仿宋_GB2312"
    assert c.force_a4 is True
    assert abs(c.margin_top_cm - 3.7) < 1e-6
    assert c.title_size == 22.0


def test_report_config():
    c = report_template()
    assert c.title_font == "黑体"
    assert c.header_enabled is True


def test_gongwen_redhead_color():
    doc = Document()
    doc.add_paragraph("原内容")
    generate_gongwen(doc)
    first = doc.paragraphs[0]
    assert first.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert len(doc.paragraphs) > 5


def test_redhead_only_red():
    doc = Document()
    generate_redhead(doc)
    assert doc.paragraphs[0].runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_report_title_bold():
    doc = Document()
    doc.add_paragraph("x")
    generate_report(doc)
    assert doc.paragraphs[0].runs[0].font.bold is True


def test_generators_covers_all_templates():
    # 每个模板种类都必须有对应的骨架生成函数
    assert set(GENERATORS.keys()) == set(TEMPLATES.keys())


def test_template_descriptions_cover_templates():
    # 描述字典应覆盖全部模板（新增模板时勿遗漏说明）
    assert set(TEMPLATE_DESCRIPTIONS.keys()) >= set(TEMPLATES.keys())
    for k in TEMPLATES:
        assert TEMPLATE_DESCRIPTIONS.get(k)


def test_generate_template_dispatch():
    # generate_template 应按 kind 统一分发到对应生成函数
    for kind in TEMPLATES:
        doc = Document()
        doc.add_paragraph("原内容")
        out = generate_template(kind, doc)
        # 生成后段落数应增加（套用了骨架）
        assert len(out.paragraphs) > 1, f"{kind} 未生成骨架"
    # 未知 kind 应原样返回
    doc2 = Document()
    assert generate_template("__nope__", doc2) is doc2


def test_config_appearance_mode():
    # 默认浅色，且不破坏 from_dict / to_dict 往返
    c = FormatterConfig()
    assert c.appearance_mode == "Light"
    d = c.to_dict()
    assert d["appearance_mode"] == "Light"
    c2 = FormatterConfig.from_dict({**d, "appearance_mode": "Dark"})
    assert c2.appearance_mode == "Dark"
    # 旧配置（无该字段）应能正常加载，回退默认浅色
    c3 = FormatterConfig.from_dict({"title_font": "黑体"})
    assert c3.appearance_mode == "Light"
