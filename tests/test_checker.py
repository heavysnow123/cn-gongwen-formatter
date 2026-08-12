import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from word_formatter.checker import check_document


def _set_eastasia(run, name):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), name)


def test_missing_font_detected(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("测试")
    _set_eastasia(r, "这个字体肯定不存在XYZ123")
    path = tmp_path / "t.docx"
    doc.save(str(path))

    rep = check_document(str(path))
    assert "这个字体肯定不存在XYZ123" in rep["missing_fonts"]
    assert rep["score"] < 100
    assert any(it["name"] == "字体完整性" and it["status"] == "fail"
               for it in rep["items"])
