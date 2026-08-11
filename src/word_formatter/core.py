"""Word Formatter 核心排版引擎。

与界面解耦，可直接被 GUI / CLI / 脚本调用。
设计原则（相对原工具有意优化）：
- 排版逻辑纯净、可测试、全部参数可配置；
- 不修改任何原始文件，所有操作在临时副本上进行；
- 跨平台：Windows 用 WPS/Word COM（见 legacy.py），其他平台用 LibreOffice 兜底。
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import logging
from typing import Callable, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import FormatterConfig

logger = logging.getLogger("word_formatter")

# ----------------------- 常量 / 正则 -----------------------
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

RE_HAS_CHINESE = re.compile(r"[\u4e00-\u9fff]")
RE_H4 = re.compile(r"^[（(]\d+[）)]")                       # (1)
RE_H2 = re.compile(r"^[（(][一二三四五六七八九十百千万零]+[）)]")  # （一）
RE_H3 = re.compile(r"^\d+[\.、]")                           # 1. 或 1、
RE_H1 = re.compile(r"^[一二三四五六七八九十百千万零]+[、．.]")     # 一、
RE_ATTACHMENT = re.compile(r"^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?")
RE_NUMERIC_TABLE_TEXT = re.compile(r"^[0-9]+([.,]\d+)?%?$")
RE_MD_HEADER = re.compile(r"^\s{0,3}#{1,6}\s+(.*)$")
RE_MD_UNORDERED = re.compile(r"^(\s*)[*+-]\s+(.*)$")
RE_MD_BOLD = re.compile(r"\*\*([^*]+)\*\*")
RE_MD_EMPH = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
RE_MD_UNDERSCORE_BOLD = re.compile(r"__([^_]+)__")
RE_MD_UNDERSCORE_EMPH = re.compile(r"(?<!_)_([^_]+)_(?!_)")
RE_MD_INLINE_CODE = re.compile(r"`([^`]+)`")
RE_MD_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
RE_MD_IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
RE_MD_HTML = re.compile(r"<[^>]+>")
RE_MD_BLOCKQUOTE = re.compile(r"^\s*>\s?(.*)$")
RE_MD_HR = re.compile(r"^\s*([-*_])(\s*\1){2,}\s*$")

ALIGN_MAP = {"left": LEFT, "center": CENTER, "right": RIGHT, "justify": JUSTIFY}

# 直接操作 lxml 元素时用到的标签（避免 python-docx 包装对象，省内存/提速）
W_P = qn("w:p")
W_TBL = qn("w:tbl")
W_PPR = qn("w:pPr")
W_JC = qn("w:jc")


# ----------------------- 低层工具 -----------------------
def set_run_font(run, cn_font: str, size: Optional[float] = None,
                 en_font: Optional[str] = None, bold: Optional[bool] = None,
                 color: Optional[RGBColor] = None):
    """设置 run 字体，含中文(eastAsia)与西文(ascii/hAnsi)分离。"""
    run.font.name = en_font or cn_font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), cn_font)
    if en_font:
        rf.set(qn("w:ascii"), en_font)
        rf.set(qn("w:hAnsi"), en_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_align(p, align):
    p.alignment = ALIGN_MAP.get(align, LEFT)


def set_outline_level(p, level: int):
    """设置大纲级别（1-9）。python-docx 1.x 无该属性，直接写 XML。"""
    if not (1 <= level <= 9):
        return
    pPr = p._p.get_or_add_pPr()
    # 移除已有的 outlineLvl
    for old in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(old)
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level - 1))
    pPr.append(ol)


def set_para_line_spacing(p, multiple: float):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = multiple


def set_para_indent(p, left_cm: float, right_cm: float, first_line_chars: int,
                    font_size: float):
    if left_cm:
        p.paragraph_format.left_indent = Cm(left_cm)
    if right_cm:
        p.paragraph_format.right_indent = Cm(right_cm)
    if first_line_chars and first_line_chars > 0:
        p.paragraph_format.first_line_indent = Pt(font_size * first_line_chars)


def _para_text(p) -> str:
    """鲁棒取段落全文（兼容 parent=None 的临时包装，避免 p.text 为 None）。"""
    try:
        return "".join(p._p.itertext())
    except Exception:
        return p.text or ""


def get_para_font_info(p):
    """返回 (eastasia_font, size_pt)，没有显式设置则为 None。"""
    name = None
    size = None
    for r in p.runs:
        if r.font.size is not None and size is None:
            size = r.font.size.pt
        rpr = r._element.find(qn("w:rPr"))
        if rpr is not None:
            rf = rpr.find(qn("w:rFonts"))
            if rf is not None:
                ea = rf.get(qn("w:eastAsia")) or rf.get(qn("w:ascii"))
                if ea and name is None:
                    name = ea
        if name and size is not None:
            break
    return name, size


def set_tbl_width(tbl, pct: int):
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(pct * 50)))
    tblW.set(qn("w:type"), "pct")


def set_tbl_borders(tbl, size_pt: float):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            borders.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(int(size_pt * 8)))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")


def set_cell_margins(tbl, top=50, left=108, bottom=50, right=108):
    tblPr = tbl._tbl.tblPr
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = mar.find(qn(f"w:{tag}"))
        if e is None:
            e = OxmlElement(f"w:{tag}")
            mar.append(e)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")


def set_row_height(row, cm: float):
    tr = row._tr
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        h = OxmlElement("w:trHeight")
        trPr.append(h)
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "atLeast")


def set_col_widths(tbl, widths_pct):
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_pct):
                continue
            tcPr = cell._tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(widths_pct[i] * 50)))
            tcW.set(qn("w:type"), "pct")


def _append_field(run, name: str):
    """往 run 里追加一个域（PAGE / NUMPAGES 等）。"""
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = name
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)


def footer_segments(cfg: FormatterConfig) -> list:
    """构造页脚片段序列：("text", str) 或 ("field", "PAGE"/"NUMPAGES")。

    标准模式（python-docx）与流式模式（XML 直写）共用，保证两套实现一致。
    """
    has_text = bool(cfg.footer_text.strip())
    want_total = cfg.page_number_total
    segs = []
    if has_text:
        segs.append(("text", cfg.footer_text.strip()))
        segs.append(("text", "    "))
    if cfg.page_number:
        if want_total:
            segs += [("text", "第 "), ("field", "PAGE"), ("text", " 页 / 共 "),
                     ("field", "NUMPAGES"), ("text", " 页")]
        elif has_text:
            segs += [("text", "第 "), ("field", "PAGE"), ("text", " 页")]
        else:
            segs.append(("field", "PAGE"))  # 仅页码：不加“第/页”装饰
    return segs


def header_font_info(cfg: FormatterConfig):
    """页眉字体信息 -> (font, size, en_font)。留空字段沿用页码设置。"""
    en = cfg.english_font if cfg.use_custom_english_font else None
    font = cfg.header_font or cfg.page_number_font or "宋体"
    size = cfg.header_size or cfg.page_number_size or 10.5
    return font, size, en


def footer_font_info(cfg: FormatterConfig):
    """页脚字体信息 -> (font, size, en_font)。"""
    en = cfg.english_font if cfg.use_custom_english_font else None
    font = cfg.page_number_font or "宋体"
    size = cfg.page_number_size or 10.5
    return font, size, en


def add_header(section, cfg: FormatterConfig):
    """插入页眉（文字 + 可选下边框线）。"""
    if not (cfg.header_enabled and cfg.header_text.strip()):
        return
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = ALIGN_MAP.get(cfg.header_align, CENTER)
    font, size, en = header_font_info(cfg)
    r = p.add_run()
    r.add_text(cfg.header_text.strip())
    set_run_font(r, font, size, en)
    if cfg.header_border:
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)


def add_footer(section, cfg: FormatterConfig):
    """插入页脚：可选附加文字 + 页码（可选“共 Y 页”）。"""
    if not (cfg.page_number or cfg.footer_text.strip()):
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = ALIGN_MAP.get(cfg.page_number_align, CENTER)
    font, size, en = footer_font_info(cfg)
    for kind, val in footer_segments(cfg):
        r = p.add_run()
        if kind == "text":
            r.add_text(val)
        else:
            _append_field(r, val)
        set_run_font(r, font, size, en)


def add_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    pPr.append(br)


# ----------------------- Markdown / 文本处理 -----------------------
def clean_markdown(text: str) -> str:
    """清理 Markdown 标记为纯文本（保留有序列表编号）。"""
    lines = text.split("\n")
    out = []
    for line in lines:
        m = RE_MD_HR.match(line)
        if m:
            continue
        m = RE_MD_HEADER.match(line)
        if m:
            out.append(m.group(1))
            continue
        m = RE_MD_BLOCKQUOTE.match(line)
        if m:
            out.append(m.group(1))
            continue
        m = RE_MD_UNORDERED.match(line)
        if m:
            out.append(m.group(2))
            continue
        out.append(line)
    text = "\n".join(out)
    text = RE_MD_IMAGE.sub("", text)
    text = RE_MD_LINK.sub(r"\1", text)
    text = RE_MD_INLINE_CODE.sub(r"\1", text)
    text = RE_MD_UNDERSCORE_BOLD.sub(r"\1", text)
    text = RE_MD_UNDERSCORE_EMPH.sub(r"\1", text)
    text = RE_MD_BOLD.sub(r"\1", text)
    text = RE_MD_EMPH.sub(r"\1", text)
    text = RE_MD_HTML.sub("", text)
    return text


def normalize_blank_lines(lines, mode: str):
    if mode == "preserve":
        return lines
    reduced = []
    prev_blank = False
    for ln in lines:
        if ln.strip() == "":
            if prev_blank:
                continue
            reduced.append(ln)
            prev_blank = True
        else:
            reduced.append(ln)
            prev_blank = False
    if mode == "remove_single":
        out = []
        n = len(reduced)
        for i, ln in enumerate(reduced):
            if (ln.strip() == "" and 0 < i < n - 1
                    and reduced[i - 1].strip() != "" and reduced[i + 1].strip() != ""):
                continue
            out.append(ln)
        return out
    return reduced  # keep_single


def normalize_punctuation(text: str) -> str:
    """保守地规范化中英文标点混用。"""
    # 三个及以上 ASCII 点 -> 省略号
    text = re.sub(r"\.{3,}", "……", text)
    # 多个 ASCII 连字符 -> 破折号
    text = re.sub(r"-{2,}", "——", text)
    # 全角空格 -> 普通空格
    text = text.replace("　", " ")
    # 中文语境下的直引号转成对直角引号（“ ”），按顺序交替开闭
    _q_state = [0]
    def _to_curly(_m):
        _q_state[0] += 1
        return "“" if _q_state[0] % 2 == 1 else "”"
    text = re.sub(r'(?<=[\u4e00-\u9fff])"(?=[\u4e00-\u9fff])', _to_curly, text)
    # 重复句末标点归并
    text = re.sub(r"([。！？])\1+", r"\1", text)
    return text


# ----------------------- 标题识别 -----------------------
def detect_heading_level(text: str):
    t = text.strip()
    if RE_H4.match(t):
        return 4
    if RE_H2.match(t):
        return 2
    if RE_H3.match(t):
        return 3
    if RE_H1.match(t):
        return 1
    return 0


def classify_paragraph(p, index, title_ids, subtitle_ids, cfg):
    """返回段落类型字符串。index 为块索引，与 detect_title_subtitle 一致。"""
    text = _para_text(p).strip()
    if not text:
        return "blank"
    if index in title_ids:
        return "title"
    if index in subtitle_ids:
        return "subtitle"
    if RE_ATTACHMENT.match(text):
        return "attachment"
    if text[:1] in ("图", "表") and p.alignment == CENTER:
        return "caption"
    lvl = detect_heading_level(text)
    if lvl:
        return f"h{lvl}"
    return "body"


def detect_title_subtitle(body, doc, source_type):
    """在 body 的直接 <w:p> 子元素中识别主/副标题。

    返回段落序号集合 (title_ids, subtitle_ids) —— 序号为「body 中 <w:p> 的计数顺序」，
    与 _format_document 的主循环计数一致。仅对开头连续居中段落建临时包装对象，
    不一次性物化全文，内存友好（大文件关键优化）。
    """
    from docx.text.paragraph import Paragraph
    title_ids = set()
    subtitle_ids = set()
    paras = [c for c in body if c.tag == W_P]
    centered = []
    for i, p_elem in enumerate(paras):
        p = Paragraph(p_elem, doc)
        if _para_text(p).strip() == "":
            if centered:
                break
            continue
        if p.alignment == CENTER:
            centered.append(i)
        else:
            break
    if not centered:
        if source_type in ("txt", "md"):
            for i, p_elem in enumerate(paras):
                p = Paragraph(p_elem, doc)
                if _para_text(p).strip():
                    title_ids.add(i)
                    return title_ids, subtitle_ids
        return title_ids, subtitle_ids

    base_name, base_size = get_para_font_info(Paragraph(paras[centered[0]], doc))
    group = "title"
    ref_name, ref_size = base_name, base_size
    for i in centered:
        name, size = get_para_font_info(Paragraph(paras[i], doc))
        same = (name == ref_name) and (size == ref_size or (name is None and ref_name is None))
        if group == "title":
            if same or (name is None and ref_name is None):
                title_ids.add(i)
            else:
                group = "subtitle"
                ref_name, ref_size = name, size
                subtitle_ids.add(i)
        else:
            if same or (name is None and ref_name is None):
                subtitle_ids.add(i)
            else:
                break
    return title_ids, subtitle_ids


# ----------------------- 格式化应用 -----------------------
def apply_para_format(p, ptype: str, cfg: FormatterConfig, source_type=None):
    en = cfg.english_font if cfg.use_custom_english_font else None

    def fmt(cn_font, size, align, indent_body=False, bold=None):
        for r in p.runs:
            set_run_font(r, cn_font, size, en, bold)
        set_para_align(p, align)
        if indent_body and cfg.first_line_indent_chars:
            set_para_indent(p, cfg.left_indent_cm, cfg.right_indent_cm,
                            cfg.first_line_indent_chars, size)
        else:
            set_para_indent(p, cfg.left_indent_cm, cfg.right_indent_cm, 0, size)

    if ptype == "title":
        fmt(cfg.title_font, cfg.title_size, "center")
        set_para_line_spacing(p, cfg.title_line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 1)
    elif ptype == "subtitle":
        fmt(cfg.subtitle_font, cfg.subtitle_size, "center")
        set_para_line_spacing(p, cfg.subtitle_line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 2)
    elif ptype == "attachment":
        fmt(cfg.attachment_font, cfg.attachment_size, "center")
    elif ptype == "caption":
        fmt(cfg.figure_caption_font, cfg.figure_caption_size, "center")
    elif ptype == "h1":
        fmt(cfg.h1_font, cfg.h1_size, "left", bold=True)
        set_para_line_spacing(p, cfg.line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 1)
    elif ptype == "h2":
        fmt(cfg.h2_font, cfg.h2_size, "left", bold=False)
        set_para_line_spacing(p, cfg.line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 2)
    elif ptype == "h3":
        fmt(cfg.body_font, cfg.body_size, "left")
        set_para_line_spacing(p, cfg.line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 3)
    elif ptype == "h4":
        fmt(cfg.body_font, cfg.body_size, "left")
        set_para_line_spacing(p, cfg.line_spacing)
        if cfg.set_outline:
            set_outline_level(p, 4)
    else:  # body
        fmt(cfg.body_font, cfg.body_size, "left", indent_body=True)
        set_para_line_spacing(p, cfg.line_spacing)


def apply_h2_inline_split(p, cfg: FormatterConfig):
    """二级标题后紧跟正文（如"（一）标题。正文..."）的段内拆分。"""
    text = _para_text(p)
    m = re.match(r"^([（(][一二三四五六七八九十百千万零]+[）)]\s*.+?[。；;])", text)
    if not m:
        return False
    split_idx = len(m.group(1))
    if split_idx >= len(text.strip()):
        return False
    en = cfg.english_font if cfg.use_custom_english_font else None
    # 按 run 累计长度切换字体
    cum = 0
    for r in p.runs:
        rlen = len(r.text)
        if cum < split_idx <= cum + rlen:
            # 该 run 跨越边界：拆成两段
            boundary = split_idx - cum
            head = r.text[:boundary]
            tail = r.text[boundary:]
            r.text = head
            set_run_font(r, cfg.h2_font, cfg.h2_size, en)
            new_r = p.add_run(tail)
            set_run_font(new_r, cfg.body_font, cfg.body_size, en)
            # 后续 run 归属正文
            idx = p.runs.index(new_r)
            for rr in p.runs[idx + 1:]:
                set_run_font(rr, cfg.body_font, cfg.body_size, en)
            p.alignment = LEFT
            set_para_line_spacing(p, cfg.line_spacing)
            return True
        elif cum + rlen <= split_idx:
            set_run_font(r, cfg.h2_font, cfg.h2_size, en)
        else:
            set_run_font(r, cfg.body_font, cfg.body_size, en)
        cum += rlen
    return True


def format_tables(tables, cfg: FormatterConfig):
    """对一组 Table 对象做表格格式化。tables 可为 doc.tables 或 [Table(elem, None)]。"""
    en = cfg.english_font if cfg.use_custom_english_font else None
    for tbl in tables:
        set_tbl_width(tbl, cfg.table_width_percent)
        if cfg.table_unified_borders:
            set_tbl_borders(tbl, cfg.table_border_size_pt)
        set_cell_margins(tbl)
        # 列宽
        ncols = len(tbl.columns)
        if cfg.table_auto_col_width and ncols:
            widths = compute_col_widths(tbl, ncols, cfg)
            set_col_widths(tbl, widths)
        for ri, row in enumerate(tbl.rows):
            set_row_height(row, cfg.table_row_height_cm)
            is_header = (ri == 0)
            for cell in row.cells:
                for p in cell.paragraphs:
                    if is_header and cfg.table_header_bold:
                        for r in p.runs:
                            set_run_font(r, cfg.table_header_font, cfg.table_size, en, bold=True)
                        if cfg.table_smart_align:
                            p.alignment = CENTER
                    else:
                        for r in p.runs:
                            set_run_font(r, cfg.table_font, cfg.table_size, en)
                        if cfg.table_smart_align:
                            txt = _para_text(p).strip()
                            if RE_NUMERIC_TABLE_TEXT.match(txt):
                                p.alignment = RIGHT
                            elif len(txt) <= cfg.table_short_text_len:
                                p.alignment = CENTER
                            else:
                                p.alignment = LEFT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = cfg.table_line_spacing


def compute_col_widths(tbl, ncols, cfg: FormatterConfig):
    """按各列最长内容比例分配宽度，并夹在 [min,max] 百分比之间。"""
    max_len = [0] * ncols
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i >= ncols:
                continue
            ln = max((len(_para_text(p).strip()) for p in cell.paragraphs), default=0)
            max_len[i] = max(max_len[i], ln)
    total = sum(max_len) or 1
    raw = [cfg.table_col_min_pct + (cfg.table_col_max_pct - cfg.table_col_min_pct) * (l / total)
           for l in max_len]
    scale = cfg.table_width_percent / (sum(raw) or 1)
    return [min(cfg.table_col_max_pct, max(cfg.table_col_min_pct, w * scale)) for w in raw]


def apply_page_setup(doc, cfg: FormatterConfig):
    for section in doc.sections:
        if cfg.force_a4:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.top_margin = Cm(cfg.margin_top_cm)
        section.bottom_margin = Cm(cfg.margin_bottom_cm)
        section.left_margin = Cm(cfg.margin_left_cm)
        section.right_margin = Cm(cfg.margin_right_cm)
        section.footer_distance = Cm(cfg.footer_distance_cm)
        add_header(section, cfg)
        add_footer(section, cfg)


# ----------------------- 主流程 -----------------------
class WordFormatter:
    def __init__(self, cfg: Optional[FormatterConfig] = None,
                 log_cb: Optional[Callable[[str, str], None]] = None):
        self.cfg = cfg or FormatterConfig()
        self.log_cb = log_cb
        self._temp_files = []

    def _log(self, msg: str, level: str = "INFO"):
        if self.log_cb:
            self.log_cb(msg, level)
        if level == "ERROR":
            logger.error(msg)
        else:
            logger.info(msg)

    def _cleanup_temp(self):
        for f in self._temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except OSError:
                pass
        self._temp_files = []

    def _make_temp(self, suffix: str) -> str:
        fd, path = tempfile.mkstemp(suffix=suffix)
        os.close(fd)
        self._temp_files.append(path)
        return path

    def _prepare_input(self, src: str, source_type: str, work_docx: str) -> Optional[str]:
        """把任意源转为可处理的 .docx 路径（临时），返回该路径或 None(跳过)。"""
        ext = os.path.splitext(src)[1].lower()
        if ext == ".docx":
            # 直接以原文件为输入，保存时另写 _formatted，绝不修改原文件；
            # 省去大文件（数百 MB）的整份磁盘拷贝。
            return src
        if ext in (".doc", ".wps"):
            # 旧格式：尝试转换
            try:
                from .legacy import convert_legacy_to_docx
                converted = convert_legacy_to_docx(src, self._log)
                if converted:
                    shutil.copy2(converted, work_docx)
                    if converted not in self._temp_files:
                        self._temp_files.append(converted)
                    return work_docx
            except Exception as e:
                self._log(f"旧格式转换失败，跳过：{src} ({e})", "WARN")
            return None
        if ext == ".txt":
            with open(src, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            self._text_to_docx(text, work_docx, is_md=False)
            return work_docx
        if ext == ".md":
            with open(src, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            self._text_to_docx(text, work_docx, is_md=True)
            return work_docx
        return None

    def _text_to_docx(self, text: str, out_path: str, is_md: bool):
        if is_md:
            text = clean_markdown(text)
        lines = normalize_blank_lines(text.split("\n"), self.cfg.blank_line_mode)
        doc = Document()
        for line in lines:
            if line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)
        doc.save(out_path)

    def format_file(self, src: str, out_dir: Optional[str] = None) -> dict:
        """格式化单个文件。返回结果字典。"""
        result = {"src": src, "output": None, "skipped": False, "error": None}
        ext = os.path.splitext(src)[1].lower()
        if ext not in (".docx", ".doc", ".wps", ".txt", ".md"):
            result["skipped"] = True
            result["error"] = "不支持的文件类型"
            return result
        source_type = {".txt": "txt", ".md": "md"}.get(ext, "docx")
        # 大文件预警（避免用户以为卡死）
        sz = 0
        try:
            sz = os.path.getsize(src)
            if sz > self.cfg.large_file_threshold_mb * 1024 * 1024:
                self._log(f"⚠ 文件较大（{sz / 1024 / 1024:.0f} MB），已启用流式处理；"
                          f"内存不足时可能较慢，请耐心等待进度提示。", "WARN")
        except OSError:
            pass
        work = None
        try:
            # 输出路径（先算好，流式模式可直接写）
            base = os.path.splitext(os.path.basename(src))[0]
            if out_dir:
                os.makedirs(out_dir, exist_ok=True)
                out_path = os.path.join(out_dir, f"{base}_formatted.docx")
            else:
                out_path = os.path.join(os.path.dirname(os.path.abspath(src)),
                                        f"{base}_formatted.docx")
            # 大文件流式模式：内存恒定，突破 python-docx 上限
            use_stream = self.cfg.streaming_mode or (
                ext == ".docx" and sz > self.cfg.large_file_threshold_mb * 1024 * 1024
            )
            if ext == ".docx" and use_stream:
                if self.cfg.streaming_mode:
                    self._log("⚡ 大文件流式模式：内存恒定，逐块读写。", "INFO")
                else:
                    self._log("⚡ 文件超过阈值，自动启用大文件流式模式（内存恒定，页边距与页眉页脚照常处理）。", "WARN")
                from .stream import stream_format_document
                stream_format_document(src, out_path, self.cfg, source_type, self._log)
                result["output"] = out_path
                self._log(f"✅ 已完成：{out_path}")
                return result

            if ext == ".docx":
                prepared = src  # 直接以原文件为输入，不拷贝
            else:
                work = self._make_temp(".docx")
                prepared = self._prepare_input(src, source_type, work)
            if prepared is None:
                result["skipped"] = True
                result["error"] = "旧格式转换失败，已跳过"
                return result
            self._log(f"正在格式化：{os.path.basename(src)}")
            doc = Document(prepared)
            self._format_document(doc, source_type)
            doc.save(out_path)
            result["output"] = out_path
            self._log(f"✅ 已完成：{out_path}")
        except Exception as e:
            result["error"] = str(e)
            self._log(f"❌ 处理失败：{src} - {e}", "ERROR")
        finally:
            self._cleanup_temp()
        return result

    def format_text(self, text: str, out_path: str, is_md: bool = False) -> dict:
        """直接排版文本（强制 A4）。"""
        result = {"output": None, "error": None}
        try:
            work = self._make_temp(".docx")
            was_force = self.cfg.force_a4
            self.cfg.force_a4 = True
            self._text_to_docx(text, work, is_md=is_md)
            doc = Document(work)
            self._format_document(doc, "txt")
            self.cfg.force_a4 = was_force
            doc.save(out_path)
            result["output"] = out_path
            self._log(f"✅ 文本排版完成：{out_path}")
        except Exception as e:
            result["error"] = str(e)
            self._log(f"❌ 文本排版失败：{e}", "ERROR")
        finally:
            self._cleanup_temp()
        return result

    def _format_document(self, doc, source_type: str):
        cfg = self.cfg
        body = doc.element.body
        title_ids, subtitle_ids = detect_title_subtitle(body, doc, source_type)

        from docx.text.paragraph import Paragraph

        # 直接遍历 body 子元素（不一次性建全文包装对象）。
        # 仅对 <w:p> 按需建临时 Paragraph 包装，处理完即丢弃，内存恒定。
        para_counter = -1
        total_paras = sum(1 for c in body if c.tag == W_P)
        step = max(1, total_paras // 10) if total_paras else 1
        done = 0
        for child in body:
            if child.tag == W_TBL:
                continue
            if child.tag != W_P:
                continue
            para_counter += 1
            p = Paragraph(child, doc)
            ptype = classify_paragraph(p, para_counter, title_ids, subtitle_ids, cfg)
            if ptype == "attachment":
                add_page_break_before(p)
                if cfg.enable_attachment_formatting:
                    apply_para_format(p, "attachment", cfg)
            elif ptype == "h2":
                if not apply_h2_inline_split(p, cfg):
                    apply_para_format(p, "h2", cfg)
            else:
                apply_para_format(p, ptype, cfg)
            done += 1
            if total_paras and done % step == 0:
                self._log(f"排版进度 {int(done / total_paras * 100)}%", "INFO")

        # 符号标准化（实验）
        if cfg.normalize_punctuation:
            for child in body:
                if child.tag != W_P:
                    continue
                p = Paragraph(child, doc)
                _pt = _para_text(p)
                if _pt.strip():
                    new = normalize_punctuation(_pt)
                    if new != _pt:
                        for r in p.runs:
                            r.text = normalize_punctuation(r.text)

        # 表格（若开启）
        if cfg.enable_table_formatting:
            format_tables(doc.tables, cfg)

        # 页面
        apply_page_setup(doc, cfg)
        if total_paras:
            self._log("排版进度 100%", "INFO")
