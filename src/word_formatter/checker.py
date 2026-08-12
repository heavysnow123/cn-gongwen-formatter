"""排版质检：检查文档字体缺失、页边距、字号、页码等指标并评分。

字体是否缺失复用 fonts.py 的系统字体检测（仅检测，不安装）。
"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn


def _run_eastasia(run):
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        return None
    return rf.get(qn("w:eastAsia"))


def check_document(docx_path: str) -> dict:
    """返回结构化质检报告。"""
    doc = Document(docx_path)

    # 1) 字体使用统计
    from .fonts import is_font_installed
    fonts_used: dict[str, int] = {}
    for p in doc.paragraphs:
        for run in p.runs:
            name = _run_eastasia(run) or run.font.name
            if name:
                fonts_used[name] = fonts_used.get(name, 0) + 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        name = _run_eastasia(run) or run.font.name
                        if name:
                            fonts_used[name] = fonts_used.get(name, 0) + 1

    missing_fonts = sorted(
        [n for n, _ in fonts_used.items() if n and not is_font_installed(n)]
    )

    # 2) 页边距（cm）
    sec = doc.sections[0]
    margins_cm = {
        "top": round(sec.top_margin.cm, 2) if sec.top_margin is not None else None,
        "bottom": round(sec.bottom_margin.cm, 2) if sec.bottom_margin is not None else None,
        "left": round(sec.left_margin.cm, 2) if sec.left_margin is not None else None,
        "right": round(sec.right_margin.cm, 2) if sec.right_margin is not None else None,
    }
    is_a4 = all(
        abs((margins_cm.get(k) or 0) - 2.54) < 0.05 for k in ("top", "bottom", "left", "right")
    )
    is_gongwen = (
        abs((margins_cm.get("top") or 0) - 3.7) < 0.1
        and abs((margins_cm.get("bottom") or 0) - 3.5) < 0.1
        and abs((margins_cm.get("left") or 0) - 2.8) < 0.1
        and abs((margins_cm.get("right") or 0) - 2.6) < 0.1
    )

    # 3) 最小字号 / 行距
    min_size = None
    has_line_spacing = False
    for p in doc.paragraphs:
        if p.paragraph_format.line_spacing is not None:
            has_line_spacing = True
        for run in p.runs:
            if run.font.size is not None:
                pt = run.font.size.pt
                if min_size is None or pt < min_size:
                    min_size = pt

    # 4) 页码（footers 是否含 PAGE 域）
    has_page_number = False
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer):
            try:
                xml = footer.paragraphs[0]._p.xml if footer.paragraphs else ""
            except Exception:
                xml = ""
            if "PAGE" in xml:
                has_page_number = True
                break

    # ---------------- 评分 ----------------
    score = 100
    items = []
    notes = []

    if missing_fonts:
        penalty = min(40, 12 * len(missing_fonts))
        score -= penalty
        items.append({
            "name": "字体完整性",
            "status": "fail",
            "detail": f"缺失 {len(missing_fonts)} 种字体：{', '.join(missing_fonts)}。"
                      f"排版将回退替代字体，版式可能不达标。",
        })
        notes.append("请在系统中安装缺失字体（仿宋_GB2312 / 楷体_GB2312 为公文必备）。")
    else:
        items.append({"name": "字体完整性", "status": "ok", "detail": "文档所用字体均已安装。"})

    if is_gongwen:
        items.append({"name": "页边距（版心）", "status": "ok",
                      "detail": "符合公文版心（上37/下35/左28/右26 mm）。"})
    elif is_a4:
        items.append({"name": "页边距（版心）", "status": "ok",
                      "detail": "常规 A4 页边距（四周 2.54 cm）。"})
    else:
        score -= 10
        items.append({"name": "页边距（版心）", "status": "warn",
                      "detail": f"当前页边距 {margins_cm} cm，非标准 A4 或公文版心。"})

    if min_size is not None and min_size < 10.5:
        score -= 10
        items.append({"name": "最小字号", "status": "warn",
                      "detail": f"文档最小字号 {min_size:g} 磅，小于小四(10.5)，可能偏小。"})
    else:
        items.append({"name": "最小字号", "status": "ok",
                      "detail": f"最小字号 {min_size:g} 磅（若已检测）。" if min_size else "未检测到显式字号。"})

    if has_line_spacing:
        items.append({"name": "行距", "status": "ok", "detail": "已设置段落行距。"})
    else:
        items.append({"name": "行距", "status": "warn", "detail": "未检测到显式行距设置。"})

    if has_page_number:
        items.append({"name": "页码", "status": "ok", "detail": "已插入页码。"})
    else:
        score -= 5
        items.append({"name": "页码", "status": "warn", "detail": "未发现页码，正式文档建议加页码。"})

    score = max(0, min(100, score))
    level = "优秀" if score >= 90 else "良好" if score >= 75 else "合格" if score >= 60 else "需改进"

    return {
        "score": score,
        "level": level,
        "items": items,
        "missing_fonts": missing_fonts,
        "fonts_used": fonts_used,
        "margins_cm": margins_cm,
        "notes": notes,
    }
