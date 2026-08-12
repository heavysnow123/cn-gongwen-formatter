"""文档模板：预设排版配置 + 标准骨架 / 红头生成。

模板 = 一组 FormatterConfig 预设参数 + 可选的文档骨架注入。
仅生成占位结构与红色发文机关标志；不伪造单位印章。
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import FormatterConfig


# ---------------- 预设配置 ----------------
def gongwen_template() -> FormatterConfig:
    """公文预设（参照 GB/T 9704-2012 常用版式）。"""
    c = FormatterConfig()
    c.title_font = "思源宋体"; c.title_size = 22.0        # 二号
    c.subtitle_font = "思源宋体"; c.subtitle_size = 16.0
    c.h1_font = "黑体"; c.h1_size = 16.0                  # 三号
    c.h2_font = "楷体_GB2312"; c.h2_size = 16.0          # 三号
    c.body_font = "仿宋_GB2312"; c.body_size = 16.0      # 三号
    c.attachment_font = "仿宋_GB2312"
    c.attachment_size = 16.0
    c.figure_caption_font = "仿宋_GB2312"; c.figure_caption_size = 14.0
    c.table_caption_font = "仿宋_GB2312"; c.table_caption_size = 14.0
    c.use_custom_english_font = True; c.english_font = "Times New Roman"
    c.title_line_spacing = 1.0; c.subtitle_line_spacing = 1.0
    c.line_spacing = 1.0
    c.left_indent_cm = 0.0; c.right_indent_cm = 0.0
    c.first_line_indent_chars = 2
    c.force_a4 = True
    c.margin_top_cm = 3.7; c.margin_bottom_cm = 3.5       # 天头37 / 地脚35
    c.margin_left_cm = 2.8; c.margin_right_cm = 2.6       # 订口28 / 翻口26
    c.page_number = True; c.page_number_align = "center"
    c.page_number_font = "仿宋_GB2312"; c.page_number_size = 14.0  # 四号
    c.footer_distance_cm = 1.75
    c.page_number_total = False
    c.header_enabled = False
    c.set_outline = True
    c.normalize_punctuation = False
    c.blank_line_mode = "remove_single"
    c.enable_attachment_formatting = True
    c.enable_table_formatting = True
    c.table_font = "仿宋_GB2312"; c.table_size = 14.0
    c.table_header_font = "黑体"; c.table_header_bold = True
    c.table_width_percent = 100; c.table_row_height_cm = 0.8
    c.table_line_spacing = 12.0; c.table_border_size_pt = 0.5
    c.table_auto_col_width = True; c.table_smart_align = True
    c.table_unified_borders = True
    return c


def report_template() -> FormatterConfig:
    """报告预设（通用单位/部门汇报版式）。"""
    c = FormatterConfig()
    c.title_font = "黑体"; c.title_size = 22.0
    c.subtitle_font = "黑体"; c.subtitle_size = 16.0
    c.h1_font = "黑体"; c.h1_size = 16.0
    c.h2_font = "黑体"; c.h2_size = 14.0
    c.body_font = "宋体"; c.body_size = 12.0             # 小四
    c.attachment_font = "黑体"; c.attachment_size = 14.0
    c.figure_caption_font = "宋体"; c.figure_caption_size = 10.5
    c.table_caption_font = "宋体"; c.table_caption_size = 10.5
    c.use_custom_english_font = True; c.english_font = "Times New Roman"
    c.title_line_spacing = 1.3; c.subtitle_line_spacing = 1.3
    c.line_spacing = 1.5
    c.first_line_indent_chars = 2
    c.force_a4 = True
    c.margin_top_cm = 2.54; c.margin_bottom_cm = 2.54
    c.margin_left_cm = 2.54; c.margin_right_cm = 2.54
    c.page_number = True; c.page_number_align = "center"
    c.page_number_font = "宋体"; c.page_number_size = 10.5
    c.footer_distance_cm = 1.75
    c.header_enabled = True; c.header_text = ""; c.header_align = "center"
    c.header_border = True
    c.set_outline = True
    c.enable_table_formatting = True
    c.table_font = "宋体"; c.table_size = 10.5
    c.table_header_font = "黑体"; c.table_header_bold = True
    return c


TEMPLATES = {
    "gongwen": ("公文", gongwen_template),
    "report": ("报告", report_template),
}


# ---------------- 骨架 / 红头生成 ----------------
def _set_run(run, cn_font, size, bold=None, color=None, en=None):
    run.font.name = en or cn_font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), cn_font)
    rf.set(qn("w:ascii"), en or cn_font)
    rf.set(qn("w:hAnsi"), en or cn_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _prepend(doc, new_p):
    """把已创建的段落移到文档最前（保持相对顺序时用 reversed 调用）。"""
    body = doc.element.body
    ref = None
    for child in body.iterchildren():
        if child.tag != qn("w:sectPr"):
            ref = child
            break
    if ref is None:
        body.append(new_p)
    else:
        ref.addprevious(new_p)


def _add_red_bottom_border(p, sz=36, color="FF0000"):
    """给段落加红色下边框（公文红线）。sz 单位：八分之一磅。"""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    bx = pBdr.find(qn("w:bottom"))
    if bx is not None:
        pBdr.remove(bx)
    pBdr.append(bottom)


def generate_gongwen(doc, org_name="（请输入发文机关名称）",
                     doc_number="（文号，如 ××〔2026〕×号）",
                     title="（公文标题）",
                     main_recv="（主送机关）：",
                     sign_org="（发文机关名称）",
                     sign_date="（成文日期）"):
    """在文档开头插入公文红头 + 标准骨架（占位文本，用户随后替换）。"""
    paras = []

    # 1) 发文机关标志（红头，红色，居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(org_name)
    _set_run(r, "思源宋体", 22.0, bold=False, color=RGBColor(0xFF, 0x00, 0x00))
    paras.append(p)

    # 2) 红色分隔线（武文线）
    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.add_run("")
    _add_red_bottom_border(pl, sz=36, color="FF0000")
    paras.append(pl)

    # 3) 文号（红线下方，左对齐于版心）
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rd = pd.add_run(doc_number)
    _set_run(rd, "仿宋_GB2312", 14.0)
    paras.append(pd)

    # 4) 标题（居中，二号）
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(title)
    _set_run(rt, "思源宋体", 22.0, bold=True)
    paras.append(pt)

    # 5) 主送机关
    pm = doc.add_paragraph()
    pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rm = pm.add_run(main_recv)
    _set_run(rm, "仿宋_GB2312", 16.0)
    paras.append(pm)

    # 6) 正文占位
    pb = doc.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rb = pb.add_run("（正文：此处填写公文主要内容，首行自动缩进 2 字符。）")
    _set_run(rb, "仿宋_GB2312", 16.0)
    paras.append(pb)

    # 7) 落款（右对齐：机关 + 日期）
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rs = ps.add_run(sign_org + "\n" + sign_date)
    _set_run(rs, "仿宋_GB2312", 16.0)
    paras.append(ps)

    for p in reversed(paras):
        _prepend(doc, p._p)
    return doc


def generate_report(doc, title="（报告标题）", subtitle="（副标题 / 单位 · 日期）"):
    """在文档开头插入报告骨架（大标题 + 副标题 + 章节占位）。"""
    paras = []

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(title)
    _set_run(rt, "黑体", 22.0, bold=True)
    paras.append(pt)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run(subtitle)
    _set_run(rs, "黑体", 16.0, bold=False)
    paras.append(ps)

    for label in ("一、章节一", "二、章节二", "三、章节三"):
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rh = ph.add_run(label)
        _set_run(rh, "黑体", 16.0, bold=True)
        paras.append(ph)
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rb = pb.add_run("（请在各级标题下填写内容。）")
        _set_run(rb, "宋体", 12.0)
        paras.append(pb)

    for p in reversed(paras):
        _prepend(doc, p._p)
    return doc


def generate_redhead(doc, org_name="（请输入发文机关名称）",
                     doc_number="（文号，如 ××〔2026〕×号）"):
    """仅生成红色发文机关标志 + 红线 + 文号占位（不含标题/正文骨架）。"""
    paras = []
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(org_name)
    _set_run(r, "思源宋体", 22.0, bold=False, color=RGBColor(0xFF, 0x00, 0x00))
    paras.append(p)

    pl = doc.add_paragraph()
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.add_run("")
    _add_red_bottom_border(pl, sz=36, color="FF0000")
    paras.append(pl)

    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rd = pd.add_run(doc_number)
    _set_run(rd, "仿宋_GB2312", 14.0)
    paras.append(pd)

    for p in reversed(paras):
        _prepend(doc, p._p)
    return doc
